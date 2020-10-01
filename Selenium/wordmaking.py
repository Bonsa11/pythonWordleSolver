# -*- coding: utf-8 -*-
"""
Created on Tue Sep 22 21:49:48 2020

@author: Sam
"""

import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'single')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


df_job = pd.read_csv('C:/Users/Sam/Documents/Python Projects/Selenium/jobscrapingdata.csv')

mydoc = docx.Document() # creating word document

mydoc.add_heading('Job Applications', 0)

for i in range(len(df_job)):
    JobTitle = df_job.iloc[i,0]
    Company = df_job.iloc[i,1]
    location = df_job.iloc[i,2]
    link = df_job.iloc[i,3]
    review = df_job.iloc[i,4]
    salary = df_job.iloc[i,5]
    description = df_job.iloc[i,6]
    DateAdded = df_job.iloc[i,7]
    
    CardTitle = mydoc.add_heading(JobTitle + ' ', 1)
    Cardcomp = mydoc.add_heading(Company + '   ' + review, 2)
    
    Cardloc = mydoc.add_heading(location, 3).add_run('     - ' + DateAdded)
    
    Carddesc = mydoc.add_paragraph()
    Carddesc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    Carddesc_format = Carddesc.paragraph_format
    Carddesc_format.space_before = Pt(18)
    Carddesc_format.space_before = Pt(18)
    
    Carddesc2 = Carddesc.add_run(description)
    
    fontdesc = Carddesc2.font
    fontdesc.name = 'Calibri'
    fontdesc.size = Pt(11)
    
    
    linkpara = mydoc.add_paragraph()
    hyperlink = add_hyperlink(linkpara, link, 'Link To Post', None, True)
    
    
    run = mydoc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

mydoc.save('C:/Users/Sam/Documents/Python Projects/Selenium/JobApplications.docx')



