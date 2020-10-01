# -*- coding: utf-
"""
Created on Mon Sep 21 16:12:12 2020

@author: Sam
"""

"""

This script will scrape indeed for jobs based on your selected criteria and output
to a word document with the the most important details including a hyperlink to the job posts

"""

#############################################################################

# importing packages

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
import pandas as pd
from datetime import date
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

#############################################################################

# defining ez functions

def small_menu_select(ID,key,times):
    search = driver.find_element_by_id(str(ID))
    search.click()
    for i in range(times): search.send_keys(str(key)) 
    search.send_keys(Keys.RETURN)
    
def findby_id_click(ID):
    search = driver.find_element_by_id(ID)
    search.click()
    
def findby_xpath_click(xpath):
    search = driver.find_element_by_xpath(xpath)
    search.click()
    
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    
    # Function taken from https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
    
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'single')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

    
#############################################################################

### YOUR VARIABLES ###

FILEPATH = 'C:/Users/Sam/Documents/Python Projects/Selenium/'
# change this to wherever you want to save the csv and word file to

PATH = 'C:\Program Files (x86)\chromedriver.exe'
# look up selenium setup tutorial to find this

WHAT = 'Graduate Data'
# what to search indeed for

WHERE = 'United Kingdom' # has to be in the UK as indeed changes address after that
# where to look

#############################################################################

### Scraping

# setting up paramaters

WHAT = WHAT.split(' ') # splits by space into list
WHERE = WHERE.split(' ')

WHAT = str('+'.join(WHAT))
WHERE = str('+'.join(WHERE))

# getting to the right place

indeedlink = 'https://www.indeed.co.uk/jobs?q='+ WHAT +'&l=' + WHERE

driver = webdriver.Chrome(PATH) # open browser

driver.get(indeedlink) # go to link

time.sleep(1)

findby_id_click('onetrust-accept-btn-handler')

time.sleep(3)

for i in range(2): 
    findby_xpath_click('/html/body/table[1]/tbody/tr/td/table/tbody/tr/td/form/table/tbody/tr[3]/td[4]/div/a')

time.sleep(2)
    
small_menu_select('fromage', 'w', 1) # uses quick menus
small_menu_select('limit', '2', 1)

findby_id_click('fj')

time.sleep(1)

findby_id_click('popover-x')

#let the driver wait 3 seconds to locate the element before exiting out
driver.implicitly_wait(3) 

# scraping the cards 

titles=[]
companies=[]
locations=[]
links =[]
reviews=[]
salaries = []
descriptions=[]
    
job_card = driver.find_elements_by_xpath('//div[contains(@class,"clickcard")]')
    
for job in job_card:
       
    #.  not all companies have review
    try:
        review = job.find_element_by_xpath('.//span[@class="ratingsContent"]').text
    except:
        review = "None"
    reviews.append(review)
   #.   not all positions have salary
    try:
        salary = job.find_element_by_xpath('.//span[@class="salaryText"]').text
    except:
        salary = "None"
    #.  tells only to look at the element       
    salaries.append(salary)
        
    try:
        location = job.find_element_by_xpath('.//span[contains(@class,"location")]').text
    except:
        location = "None"
    #.  tells only to look at the element       
    locations.append(location)
        
    try:
        title  = job.find_element_by_xpath('.//h2[@class="title"]//a').text
    except:
        title = job.find_element_by_xpath('.//h2[@class="title"]//a').get_attribute(name="title")
    titles.append(title)
    links.append(job.find_element_by_xpath('.//h2[@class="title"]//a').get_attribute(name="href"))
    companies.append(job.find_element_by_xpath('.//span[@class="company"]').text)


descriptions=[]
for link in links:
    
    driver.get(link)
    jd = driver.find_element_by_xpath('//div[@id="jobDescriptionText"]').text
    descriptions.append(jd)
    
driver.close()

df_new = pd.DataFrame()
df_new['Title']=titles
df_new['Company']=companies
df_new['Location']=locations
df_new['Link']=links
df_new['Review']=reviews
df_new['Salary']=salaries
df_new['Description']=descriptions
df_new['Date Added'] = date.today()
df_new['Applied'] = 0

for i in range(20):
    df_new['Title'][i] = df_new['Title'][i].upper()
    df_new['Company'][i] = df_new['Company'][i].upper()
    df_new['Location'][i] = df_new['Location'][i].upper()
    df_new['Description'][i] = df_new['Description'][i].lower()

df_job = pd.read_csv(FILEPATH + 'jobscrapingdata.csv') # read in current

df_full = df_job.append(df_new, ignore_index=True)
df_full = df_full.drop_duplicates(subset=['Description']) # removes duplicate listings
# no primary key, links changed so didn't eliminate  

df_full.to_csv(FILEPATH + 'jobscrapingdata.csv', index = False)

#############################################################################

### writing word document

df_job = pd.read_csv(FILEPATH + 'jobscrapingdata.csv')

mydoc = docx.Document() # creating word document

mydoc.add_heading('Job Applications', 0)

for i in range(len(df_job)):
    JobTitle = df_job.iloc[i,0]
    Company = df_job.iloc[i,1]
    location = df_job.iloc[i,2]
    link = df_job.iloc[i,3]
    review = df_job.iloc[i,4]
    salary = df_job.iloc[i,5]
    description = df_job.iloc[i,6]
    DateAdded = df_job.iloc[i,7]
    
    CardTitle = mydoc.add_heading('Job Title:  ' + JobTitle + ' ', 1)
    Cardcomp = mydoc.add_heading( 'Company:  ' + Company + '   ' + '      Review:  ' + review, 2)
    
    Cardloc = mydoc.add_heading('Location:  ' + location, 3).add_run('      Date:  ' + DateAdded)
    
    Carddesc = mydoc.add_paragraph()
    Carddesc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    Carddesc_format = Carddesc.paragraph_format
    Carddesc_format.space_before = Pt(18)
    Carddesc_format.space_before = Pt(18)
    
    Carddesc2 = Carddesc.add_run(description)
    
    fontdesc = Carddesc2.font
    fontdesc.name = 'Calibri'
    fontdesc.size = Pt(11)
    
    
    linkpara = mydoc.add_paragraph()
    hyperlink = add_hyperlink(linkpara, link, 'Link To Post', None, True)
    
    
    run = mydoc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

mydoc.save(FILEPATH + 'JobApplications.docx')
